from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Configuração de página ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.1)
section.bottom_margin = Inches(1.1)

# ── Paleta de cores ─────────────────────────────────────────────────────────
ROXO_ESCURO  = RGBColor(0x4A, 0x00, 0x8F)   # #4A008F
ROXO_MEDIO   = RGBColor(0x7B, 0x2D, 0xBF)   # #7B2DBF
ROSA_DESTAQUE= RGBColor(0xE0, 0x3E, 0x9C)   # #E03E9C
CINZA_ESCURO = RGBColor(0x23, 0x23, 0x23)   # #232323
CINZA_TEXTO  = RGBColor(0x44, 0x44, 0x44)   # #444444
BRANCO       = RGBColor(0xFF, 0xFF, 0xFF)

def set_run_color(run, color):
    run.font.color.rgb = color

def add_heading(doc, text, level=1, color=None, bold=True, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        set_run_color(run, color)
    if size:
        run.font.size = Pt(size)
    else:
        sizes = {1: 22, 2: 16, 3: 13}
        run.font.size = Pt(sizes.get(level, 12))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_body(doc, text, indent=False, color=None, size=11):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        set_run_color(run, color)
    else:
        set_run_color(run, CINZA_TEXTO)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(11)
        set_run_color(r1, ROXO_ESCURO)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        set_run_color(r2, CINZA_TEXTO)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        set_run_color(run, CINZA_TEXTO)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 72)
    run.font.size = Pt(9)
    set_run_color(run, RGBColor(0xCC, 0xCC, 0xCC))
    return p

def add_feature_block(doc, titulo, descricao, beneficio):
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(10)
    p_titulo.paragraph_format.space_after  = Pt(2)
    p_titulo.paragraph_format.left_indent  = Inches(0.2)
    r = p_titulo.add_run(f"  {titulo}")
    r.bold = True
    r.font.size = Pt(12)
    set_run_color(r, ROXO_MEDIO)

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.left_indent = Inches(0.5)
    p_desc.paragraph_format.space_after = Pt(2)
    rd = p_desc.add_run(descricao)
    rd.font.size = Pt(11)
    set_run_color(rd, CINZA_TEXTO)

    p_ben = doc.add_paragraph()
    p_ben.paragraph_format.left_indent = Inches(0.5)
    p_ben.paragraph_format.space_after = Pt(6)
    rb1 = p_ben.add_run("Benefício: ")
    rb1.bold = True
    rb1.font.size = Pt(11)
    set_run_color(rb1, ROSA_DESTAQUE)
    rb2 = p_ben.add_run(beneficio)
    rb2.font.size = Pt(11)
    set_run_color(rb2, CINZA_TEXTO)


# ════════════════════════════════════════════════════════════════════════════
#  CAPA
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_logo = p_logo.add_run("PONTO DA DANÇA")
r_logo.bold = True
r_logo.font.size = Pt(36)
set_run_color(r_logo, ROXO_ESCURO)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("Sistema de Gestão para Academias de Dança")
r_sub.font.size = Pt(18)
r_sub.bold = False
set_run_color(r_sub, ROXO_MEDIO)
p_sub.paragraph_format.space_before = Pt(6)

doc.add_paragraph()

p_tagline = doc.add_paragraph()
p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_tag = p_tagline.add_run("Mais organização. Menos papel. Mais tempo para dançar.")
r_tag.italic = True
r_tag.font.size = Pt(13)
set_run_color(r_tag, ROSA_DESTAQUE)
p_tagline.paragraph_format.space_before = Pt(4)

doc.add_paragraph()
doc.add_paragraph()

p_divs = doc.add_paragraph()
p_divs.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_divs = p_divs.add_run("━" * 50)
r_divs.font.size = Pt(10)
set_run_color(r_divs, RGBColor(0xBB, 0xBB, 0xBB))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. O PROBLEMA ATUAL
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "O Problema Atual nas Academias de Dança", level=1, color=ROXO_ESCURO, size=20)
add_divider(doc)

add_body(doc,
    "Gerir uma academia de dança vai muito além de ensinar passos. "
    "Todo dia surgem desafios operacionais que consomem horas do dono e da equipe — "
    "e que, se mal resolvidos, custam alunos, dinheiro e credibilidade.")

add_heading(doc, "Dores mais comuns dos donos de escola", level=2, color=ROXO_MEDIO, size=14)

dores = [
    ("Controle manual e planilhas desorganizadas:",
     "frequências anotadas em papel, listas espalhadas, dados perdidos."),
    ("Comunicação ineficiente:",
     "avisos por grupos de WhatsApp misturados, informações que não chegam."),
    ("Chamada demorada e trabalhosa:",
     "professor perde tempo em sala de aula com papel em vez de dançar."),
    ("Gestão de reposições sem controle:",
     "alunos ligam para saber se há vaga; recepção não sabe ao certo."),
    ("Bolsistas sem acompanhamento:",
     "difícil saber se o aluno bolsista está correspondendo às expectativas."),
    ("Aulas particulares agendadas no \"boca a boca\":",
     "conflitos de horário, esquecimentos e falta de histórico."),
    ("Sem presença digital:",
     "escola sem página na internet perde alunos para quem está no Google."),
    ("Acesso desigual à informação:",
     "o aluno não sabe o horário da próxima aula; o professor não vê a lista da turma."),
]

for bold_text, rest in dores:
    add_bullet(doc, rest, bold_prefix=bold_text)

add_body(doc,
    "\nO resultado? Horas desperdiçadas, erros evitáveis, insatisfação de alunos e professores — "
    "e um dono de escola que trabalha dobrado sem ver a escola crescer.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  2. A SOLUÇÃO
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "A Solução: Ponto da Dança", level=1, color=ROXO_ESCURO, size=20)
add_divider(doc)

add_body(doc,
    "O Ponto da Dança é um sistema de gestão criado do zero para academias de dança. "
    "Não é uma planilha adaptada, não é um sistema genérico de academia fitness — "
    "é uma ferramenta pensada nos detalhes do dia a dia de quem vive a dança.")

add_body(doc,
    "Funciona direto no navegador do celular ou computador, sem precisar instalar nada. "
    "Cada pessoa da sua escola — gerente, recepção, professor e aluno — "
    "tem acesso ao que precisa, de onde estiver, com total segurança.")

add_heading(doc, "Em uma frase:", level=2, color=ROXO_MEDIO, size=13)

p_frase = doc.add_paragraph()
p_frase.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_frase.paragraph_format.space_before = Pt(6)
p_frase.paragraph_format.space_after  = Pt(6)
p_frase.paragraph_format.left_indent  = Inches(0.5)
p_frase.paragraph_format.right_indent = Inches(0.5)
r_frase = p_frase.add_run(
    '"O Ponto da Dança centraliza toda a gestão da sua escola em um só lugar, '
    'acessível pelo celular, para que você gaste energia no que importa: a dança."')
r_frase.italic = True
r_frase.font.size = Pt(13)
set_run_color(r_frase, ROXO_MEDIO)

add_heading(doc, "Quem usa o sistema:", level=2, color=ROXO_MEDIO, size=13)

perfis = [
    ("Gerente (dono da escola):", "visão completa, controle total de turmas, professores, alunos e bolsistas."),
    ("Recepção:", "gestão operacional do dia a dia — matrículas, reposições, avisos e atendimento."),
    ("Professor:", "acesso às próprias turmas, realização de chamada pelo celular e gestão de aulas particulares."),
    ("Aluno:", "visualização do próprio perfil, turmas, avisos e solicitação de reposição."),
]

for bold_text, rest in perfis:
    add_bullet(doc, rest, bold_prefix=bold_text)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  3. FUNCIONALIDADES IMPLEMENTADAS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Funcionalidades Disponíveis Hoje", level=1, color=ROXO_ESCURO, size=20)
add_divider(doc)

add_body(doc,
    "O sistema já está em funcionamento com todas as funcionalidades essenciais "
    "para o dia a dia da sua escola. Veja o que está disponível agora:")

features = [
    (
        "Gestão de Turmas",
        "Crie e organize todas as turmas da escola: ballet, forró, jazz, contemporâneo, e mais. "
        "Defina dias, horários, modalidade e o professor responsável por cada turma. "
        "Visualize a grade horária completa da escola em um único quadro.",
        "Fim das dúvidas sobre \"qual professor está em qual turma?\" — "
        "tudo centralizado e consultável por qualquer membro da equipe."
    ),
    (
        "Sistema de Chamada Digital",
        "O professor abre o celular, seleciona a turma e a data, e faz a chamada na hora — "
        "marcando presença ou falta para cada aluno. O histórico fica salvo automaticamente.",
        "Professores saem da sala de aula; alunos têm frequência registrada corretamente; "
        "donos acompanham presença sem precisar perguntar."
    ),
    (
        "Aulas Particulares",
        "Professores cadastram seus horários disponíveis para aulas particulares. "
        "A recepção ou o próprio aluno agenda a aula, e o histórico fica registrado no sistema.",
        "Sem conflito de horários, sem agendamentos esquecidos, com histórico completo de cada aluno."
    ),
    (
        "Sistema de Bolsistas",
        "Cadastro completo de alunos bolsistas com tipo de bolsa. "
        "Quadro de desempenho para o gerente acompanhar se o bolsista está cumprindo suas obrigações. "
        "Gestão diferenciada para quem tem benefícios.",
        "Controle profissional dos bolsistas — raridade no mercado de sistemas para dança. "
        "Decisões baseadas em dados, não em memória."
    ),
    (
        "Avisos e Comunicação",
        "Envie avisos para uma turma específica ou para toda a escola. "
        "Professores e alunos veem os avisos ao entrar no sistema. "
        "Histórico completo de comunicações.",
        "Chega de mensagens perdidas no WhatsApp. Avisos importantes chegam a quem precisa ver, "
        "e ficam registrados."
    ),
    (
        "Reposição de Aulas",
        "Alunos solicitam reposição de aulas perdidas pelo sistema. "
        "A recepção gerencia as vagas disponíveis e confirma ou redireciona as solicitações.",
        "Processo organizado, sem ligações desnecessárias, "
        "sem vagas sendo ocupadas duas vezes."
    ),
    (
        "Página Pública da Escola",
        "Sua escola ganha uma página na internet com as modalidades oferecidas, "
        "informações de contato e localização — tudo gerenciado pelo próprio sistema.",
        "Presença digital profissional sem precisar contratar web designer. "
        "Novos alunos encontram sua escola no Google."
    ),
    (
        "Perfis e Controle de Acesso",
        "Cada usuário vê apenas o que é relevante para o seu papel. "
        "Gerente, recepção, professor e aluno têm telas e permissões diferentes. "
        "Acesso seguro com login e senha.",
        "Informações confidenciais protegidas. Cada pessoa usa o sistema com simplicidade "
        "— sem se perder em menus que não são para ela."
    ),
]

for titulo, descricao, beneficio in features:
    add_feature_block(doc, titulo, descricao, beneficio)
    add_divider(doc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  4. FUNCIONALIDADES FUTURAS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Roadmap — O Que Vem a Seguir", level=1, color=ROXO_ESCURO, size=20)
add_divider(doc)

add_body(doc,
    "O Ponto da Dança é um sistema em constante evolução. "
    "As próximas funcionalidades já estão mapeadas e serão entregues em fases:")

# Fase 1
add_heading(doc, "Fase 1 — Em Desenvolvimento", level=2, color=ROXO_MEDIO, size=14)

fase1 = [
    ("Dashboard da Recepção:",
     "visão geral completa do dia — alunos esperados, reposições, avisos e turmas em andamento."),
    ("Ações avançadas no quadro de bolsistas:",
     "relatórios de desempenho, alertas automáticos e exportação."),
    ("Quadro de turmas para o aluno:",
     "o próprio aluno visualiza o calendário das suas turmas."),
]
for b, t in fase1:
    add_bullet(doc, t, bold_prefix=b)

# Fase 2
add_heading(doc, "Fase 2 — Planejado", level=2, color=ROXO_MEDIO, size=14)

fase2 = [
    ("Módulo Financeiro:",
     "controle de mensalidades, registro de pagamentos e alerta de inadimplência."),
    ("Eventos e Ingressos:",
     "gestão de shows e recitais com venda de ingressos online para as famílias."),
    ("Aula Experimental:",
     "agendamento simplificado para novos alunos conhecerem a escola."),
    ("Notificações Push:",
     "alertas automáticos no celular — lembretes de aula, aviso de reposição confirmada, etc."),
]
for b, t in fase2:
    add_bullet(doc, t, bold_prefix=b)

# Fase 3
add_heading(doc, "Fase 3 — Futuro", level=2, color=ROXO_MEDIO, size=14)

fase3 = [
    ("Dashboard de Métricas:",
     "relatórios de crescimento, frequência média, taxa de retenção de alunos."),
    ("Automações com Inteligência Artificial:",
     "sugestões de reposição, previsão de evasão, otimização de grade horária."),
    ("Integração com WhatsApp:",
     "envio automático de avisos, cobranças e confirmações direto no WhatsApp do aluno."),
]
for b, t in fase3:
    add_bullet(doc, t, bold_prefix=b)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  5. DIFERENCIAIS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Por Que Escolher o Ponto da Dança?", level=1, color=ROXO_ESCURO, size=20)
add_divider(doc)

diferenciais = [
    ("Feito exclusivamente para dança:",
     "não é um sistema genérico adaptado. Cada funcionalidade foi pensada para a realidade "
     "de academias de dança — modalidades, chamada por turma, bolsistas, reposições."),
    ("Funciona como app no celular:",
     "tecnologia PWA — o professor abre no celular sem instalar nada. "
     "Funciona em iPhone, Android e computador com a mesma qualidade."),
    ("Controle de bolsistas:",
     "pouquíssimos sistemas do mercado oferecem gestão de bolsistas. "
     "No Ponto da Dança, isso é nativo e completo."),
    ("Multi-perfil com acesso seguro:",
     "cada pessoa vê só o que precisa. Simples para o aluno, completo para o gerente."),
    ("Acesso remoto para todos:",
     "professor faz chamada do celular em sala de aula; "
     "aluno consulta o horário de casa; dono monitora tudo de onde estiver."),
    ("Sistema em constante evolução:",
     "módulo financeiro, eventos, notificações e inteligência artificial já estão no roadmap. "
     "Quem adota agora acompanha cada nova entrega."),
    ("Suporte próximo:",
     "desenvolvimento nacional, suporte em português, adaptação às necessidades da escola."),
]

for b, t in diferenciais:
    add_bullet(doc, t, bold_prefix=b)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  6. BENEFÍCIOS DIRETOS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Benefícios Diretos para Sua Escola", level=1, color=ROXO_ESCURO, size=20)
add_divider(doc)

add_body(doc,
    "Adotar o Ponto da Dança não é apenas trocar planilhas por um sistema. "
    "É uma mudança na forma como a escola funciona — com impacto direto no tempo, "
    "no dinheiro e na experiência de alunos e professores.")

add_heading(doc, "Tempo economizado", level=2, color=ROXO_MEDIO, size=14)

tempo = [
    "Professor faz chamada em 1 minuto, direto do celular, sem papel.",
    "Recepção consulta informações de qualquer aluno em segundos.",
    "Dono não precisa ligar para professores para saber o que aconteceu na aula.",
    "Avisos chegam a todos de uma vez — sem precisar enviar mensagem individual.",
    "Solicitações de reposição chegam organizadas, sem ligações repetidas.",
]
for t in tempo:
    add_bullet(doc, t)

add_heading(doc, "Erros evitados", level=2, color=ROXO_MEDIO, size=14)

erros = [
    "Frequência registrada automaticamente, sem risco de perda de dados.",
    "Agendamentos de aulas particulares sem conflito de horário.",
    "Vagas de reposição controladas — sem a mesma vaga sendo ocupada duas vezes.",
    "Avisos chegam ao público certo — sem confusão entre turmas diferentes.",
]
for e in erros:
    add_bullet(doc, e)

add_heading(doc, "Profissionalismo e experiência do aluno", level=2, color=ROXO_MEDIO, size=14)

prof = [
    "Alunos sentem que a escola é organizada e moderna — e isso fideliza.",
    "Página pública transmite credibilidade para quem está pesquisando escolas.",
    "Comunicação clara e registrada gera confiança nos pais e responsáveis.",
    "Professores mais preparados e focados na aula, sem burocracia.",
]
for p in prof:
    add_bullet(doc, p)

add_heading(doc, "Crescimento sustentável", level=2, color=ROXO_MEDIO, size=14)

crescimento = [
    "Com o financeiro vindo no roadmap, sua escola terá controle total de inadimplência.",
    "Eventos e venda de ingressos online ampliam a receita da escola.",
    "Dados e métricas futuras ajudarão a tomar decisões baseadas em fatos, não em intuição.",
]
for c in crescimento:
    add_bullet(doc, c)

# Fechamento
doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

p_fecha = doc.add_paragraph()
p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fecha = p_fecha.add_run(
    "O Ponto da Dança foi criado por quem entende de dança e de tecnologia.\n"
    "Sua escola merece um sistema que acompanha o seu ritmo."
)
r_fecha.italic = True
r_fecha.bold = True
r_fecha.font.size = Pt(13)
set_run_color(r_fecha, ROXO_ESCURO)

doc.add_paragraph()

p_cta = doc.add_paragraph()
p_cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cta = p_cta.add_run("Entre em contato e agende uma demonstração.")
r_cta.bold = True
r_cta.font.size = Pt(14)
set_run_color(r_cta, ROSA_DESTAQUE)

# ── Salvar ──────────────────────────────────────────────────────────────────
output_path = r"c:/Users/gdmat/source/repos/Rascunho/documentacao_ponto_da_danca.docx"
doc.save(output_path)
print(f"Documento salvo em: {output_path}")
